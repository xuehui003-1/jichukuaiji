#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《03-E1_开学点单表_学生填写版》Word 文档（可打印 A4）。
正面：引导语 + 基本信息 + 6 题；背面：暂停点记录区（第一课练记录册用）。
版本号机制同规划文档：改 VERSION / DATE / REVISION_HISTORY 后重跑，旧版保留。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============ 每次修订只需要改这里 ============
VERSION = "v1.0"
DATE = "2026-09-03"
REVISION_HISTORY = [
    ("v1.0", "2026-09-03",
     "新建：配套《03_项目一_第1讲_开学第一课_规划文档_v2.0》。正面 6 题＋基本信息；背面暂停点记录区。"),
]
# =============================================

BASE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE)
OUT = os.path.join(ROOT, "03-E1_开学点单表_学生填写版_{}_{}.docx".format(VERSION, DATE.replace("-", "")))

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2)
sec.top_margin = sec.bottom_margin = Cm(2)

GRAY = (89, 89, 89)
BLUE = (31, 78, 121)


def set_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr()
    rfonts = run._element.rPr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(text="", size=10.5, bold=False, align=None, name="宋体",
         color=None, space_after=4, space_before=0, line=1.3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, name=name, size=size, bold=bold, color=color)
    return p


def h1(text):
    para(text, size=14, bold=True, name="黑体", color=BLUE, space_before=8, space_after=4)


def blank_line(lines=1):
    for _ in range(lines):
        para("＿" * 46, size=10.5, color=GRAY, space_after=2, line=1.6)


def add_footer():
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("03-E1 开学点单表 · {} · {} ｜ 你写什么，这学期课上就出现什么".format(VERSION, DATE))
    set_font(r, size=9, color=GRAY)


add_footer()

# ================= 正面 =================
para("《财务机器人应用与开发》开学点单表", size=20, bold=True, name="黑体",
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=4)
para("你在这张纸上写什么，这学期课上就出现什么", size=12, name="黑体",
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
para("这不是考试，没有标准答案。学号一定要填对。看不懂的题按字面选就行。当堂填、当堂交。",
     size=10.5, space_after=8)
para("班级：＿＿＿＿＿＿＿＿　　姓名：＿＿＿＿＿＿＿＿　　学号：＿＿＿＿＿＿＿＿", size=11, space_after=2)
para("我的互审搭档（同桌）签名：＿＿＿＿＿＿＿＿", size=11, space_after=8)

h1("第 1 题　给我们的 AI 员工起个名字")
para("要求：一眼能看出它干什么（比如「小财」「票票」）。名字起得好，全班投票采用。", size=10.5, space_after=2)
para("我起的名字：", size=11, bold=True, space_after=1); blank_line(1)
para("备选：", size=11, bold=True, space_after=8); blank_line(1)

h1("第 2 题　它的形象风格（打一个钩）")
para("□ 可爱　　□ 酷　　□ 科技感　　□ 稳重　　□ 其他：＿＿＿＿＿＿", size=12, space_after=8)

h1("第 3 题　你最想让机器人帮你干的活（打钩，可多选）")
para("□ 抄表格　　□ 开发票　　□ 对账　　□ 催款　　□ 整理文件　　□ 回消息　　□ 其他：＿＿＿＿＿＿",
     size=12, space_after=8)

h1("第 4 题　你平时用过哪些 AI（打钩，可多选）")
para("□ 刷脸/指纹支付　　□ 语音助手（小爱、小度）　　□ AI 写文案　　□ AI 做图　　□ AI 聊天　　□ 都没用过",
     size=12, space_after=8)

h1("第 5 题　你觉得「给机器人写规矩」难不难？")
para("□ 不难　　□ 有点难　　□ 很难", size=12, space_after=2)
para("一句话说说为什么：", size=11, bold=True, space_after=8); blank_line(1)

h1("第 6 题　愿不愿意被画进教学漫画 / 彩蛋？")
para("□ 愿意（署名）　　□ 愿意（匿名）　　□ 不要", size=12, space_after=4)

# ================= 背面 =================
doc.add_page_break()
para("背面　·　暂停点记录区（第一课用）", size=14, bold=True, name="黑体", color=BLUE,
     space_after=6, space_before=6)
para("课堂上有两次暂停，先写你的猜测 / 判断，再听讲。猜错不扣分；猜错的地方，正是今天要学的地方。",
     size=10.5, space_after=10)

h1("【暂停点①】给扫地机器人加一条新规矩，让它更聪明")
para("我的答案：", size=11, bold=True, space_after=1); blank_line(3)

h1("【暂停点②】刚才两组同学写的「命令」，哪句更好？差在哪？")
para("我的判断：", size=11, bold=True, space_after=1); blank_line(3)

para("—— 下课交给老师，下一节课我们就用它 ——", size=10, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)

doc.save(OUT)
print("已生成:", OUT)
