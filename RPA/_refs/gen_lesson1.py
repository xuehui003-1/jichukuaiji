#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《03_项目一_第1讲_开学第一课_规划文档》Word 文档。
v2.0 生活化改版：不涉及财务共享中心等深奥专业概念；全用生活例子引入；
AI 智能体体验案例（两版命令对比→起名建档→学生当指挥）建立信心；
暂停点 2 处（练记录册）+ 抽人 4 处（熟悉随机点名）；分组已定方案 B（同桌互审）。
版本号机制同规划文档：改 VERSION / DATE / REVISION_HISTORY 后重跑，旧版保留。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============ 每次修订只需要改这里 ============
VERSION = "v2.1"
DATE = "2026-09-03"
REVISION_HISTORY = [
    ("v2.1", "2026-09-03",
     "①「开学点单表」更名并重做为《03-E1_开学调查问卷_学生填写版_v1.0》，与记录册彻底分开；"
     "②暂停点改为「写在记录册第 1/2 页」（记录册 05_ 将在第一讲课件定稿后配套制作，与课件一一对应，"
     "印制前先用草稿纸）；③配套新增第一节课课件 04_（HTML）；④RPA 工作区只保留每份文档最新版。"),
    ("v2.0", "2026-09-03",
     "按反馈全面生活化改版：①去除「财务共享中心」等深奥专业概念；②全部改用生活例子引入"
     "（洗衣机/扫地机器人/自动售货机/手机）；③新增 AI 智能体体验案例——两版命令对比→起名建档→"
     "学生当指挥，建立信心；④删去理论讲解，改用生活化讲法；⑤暂停点 2 处＋抽人 4 处，"
     "训练记录册与随机点名习惯；⑥分组定稿方案 B 同桌互审；⑦点单表成品单独生成 03-E1。"),
    ("v1.0", "2026-08-31",
     "新建：第一节课（项目一第1讲·开学第一课）规划文档。90 分钟分钟级课堂流程；"
     "分组两套方案对比（固定小组 / 同桌互审）与推荐；开学点单表 6 题设计；"
     "AI 点睛「财务部来了新同事」落位；趣味任务卡 A1/A2。"),
]
# =============================================

BASE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE)
OUT = os.path.join(ROOT, "03_项目一_第1讲_开学第一课_规划文档_{}_{}.docx".format(VERSION, DATE.replace("-", "")))

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2)
sec.top_margin = sec.bottom_margin = Cm(2)

GRAY = (89, 89, 89)
BLUE = (31, 78, 121)


def set_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr()
    rfonts = run._element.rPr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(text="", size=10.5, bold=False, align=None, name="宋体",
         color=None, space_after=4, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.3
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, name=name, size=size, bold=bold, color=color)
    return p


def h1(text):
    para(text, size=15, bold=True, name="黑体", color=BLUE, space_before=12, space_after=6)


def h2(text):
    para(text, size=12, bold=True, name="黑体", color=(0, 0, 0), space_before=8, space_after=4)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def make_table(headers, rows, widths_cm, font_size=9.5, header_bg="D9E2F3"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(htxt)
        set_font(r, size=font_size, bold=True)
        shade_cell(hdr[i], header_bg)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(v))
            set_font(r, size=font_size)
    for i, w in enumerate(widths_cm):
        for row in t.rows:
            row.cells[i].width = Cm(w)
    return t


def add_footer():
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("第 ")
    set_font(r1, size=9, color=GRAY)
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(it); run._r.append(f2)
    set_font(run, size=9, color=GRAY)
    r2 = p.add_run(" 页 · 共 ")
    set_font(r2, size=9, color=GRAY)
    run2 = p.add_run()
    g1 = OxmlElement("w:fldChar"); g1.set(qn("w:fldCharType"), "begin")
    it2 = OxmlElement("w:instrText"); it2.set(qn("xml:space"), "preserve"); it2.text = "NUMPAGES"
    g2 = OxmlElement("w:fldChar"); g2.set(qn("w:fldCharType"), "end")
    run2._r.append(g1); run2._r.append(it2); run2._r.append(g2)
    set_font(run2, size=9, color=GRAY)
    r3 = p.add_run(" 页")
    set_font(r3, size=9, color=GRAY)


add_footer()

# ================= 标题区 =================
para("项目一　第 1 讲　开学第一课（讲课规划文档）", size=20, bold=True, name="黑体",
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, space_before=6)
para("财务机器人应用与开发 · 第 1 周 · 90 分钟（2 课时连上）", size=12, name="黑体",
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("文档编号 03 ｜ 版本 {} ｜ 生成日期 {} ｜ 配套《00_财务机器人课程教学规划_v3.4》｜ 性质：讲课规划（第一课专属版·生活化改版）".format(VERSION, DATE),
     size=10.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# ================= 0 使用说明 =================
h1("0　使用说明（这份怎么用）")
for s in [
    "① 这是开学第一课的完整规划，90 分钟课堂流程精确到分钟，照着走就行。",
    "② 本课全程不出现「财务共享中心」这类专业概念，也不做理论讲解——所有引入都用学生家里的东西"
    "（洗衣机、扫地机器人、自动售货机、手机），专业术语这学期从项目二开始逐个「翻译」进来。",
    "③ 第一课不碰软件：智多星下周（项目二）才安装；本课只需要投影＋一台能联网的电脑"
    "（大模型网页版，豆包 / DeepSeek / 文心一言任一即可）。",
    "④ 本课定位一句话：不是教多少知识，而是让学生「下周还想来」＋「这事我能行」——期待感与信心第一。",
    "⑤ 配套文件：①《03-E1_开学调查问卷_学生填写版》已单独生成（参照基础会计课「学生入口.html」结构，"
    "2 页可打印），课前打印；②《04_项目一_第1讲_开学第一课_课件_v1.0.html》为本讲课件（含教师提示面板）；"
    "③记录册《我的机器人实验记录》（编号 05_）将在本讲课件定稿后配套制作，暂停点与课件一一对应，印制前先用草稿纸；"
    "随机点名用基础会计课现成的 90-B2 工具。",
]:
    para(s, size=10.5, space_after=3)

# ================= 一、定位与目标 =================
h1("一、本讲定位与目标")
make_table(
    ["维度", "目标"],
    [
        ("知识", "知道这门课是「学给机器人派活、定规矩」；知道 AI 智能体听话、但要把话说清楚"),
        ("技能", "会写一句「清楚的命令」让 AI 干一件小事（同桌合作）"),
        ("情感（核心）", "建立信心——“我也能指挥 AI / 机器人”；期待感拉满——“下周还想来”"),
    ],
    [2.6, 14.4], font_size=10)
para("第一课的「三不」：不讲深（零专业术语）、不动软件（下周才装智多星）、不考试（只玩）。",
     size=10.5, bold=True, space_before=4)

# ================= 二、设计理念 =================
h1("二、设计理念（本次改版的五条原则）")
for s in [
    "① 生活入口：所有例子来自学生每天见的东西——洗衣机、扫地机器人、自动售货机、手机。进门零门槛。",
    "② 少理论多体验：不讲 AI 发展史、不讲概念；用「当场指挥 AI 智能体」让学生亲身体会——AI 你让它做什么它就做什么，"
    "但要把话说清楚，这和指挥机器人一模一样。",
    "③ 建立信心：学生没学过这类课，最容易怕「背公式、背代码」。本课反复传递：这门课不靠背，靠想清楚，"
    "想清楚一件事怎么干，机器人就替你干。",
    "④ 立规矩：暂停点 2 处（练记录册：先写自己的，再听讲）；抽人 4 处（用 90-B2 随机点名，让学生熟悉这个环节）。",
    "⑤ 专业名词零门槛：凡是要用到术语的地方，先用生活说法顶替——顺序＝先…再…；条件＝如果…就…；"
    "循环＝一直做/做几遍；变量＝机器人装东西的盒子。正式术语从项目二、三逐个换回来。",
]:
    para(s, size=10.5, space_after=3)

# ================= 三、课前准备 =================
h1("三、课前准备")
for s in [
    "教师：大模型登录投影（豆包 / DeepSeek / 文心一言网页版）；预录屏兜底 2 段（①「两版命令对比」"
    "②「起名建档」生成过程，防无网）；开学调查问卷 03-E1 打印 60 份（2 页）；90-B2 随机点名工具；"
    "积分规则一页 PPT；开场视频（生活机器人混剪：扫地机、洗衣机、自动售货机、手机自动备份）。",
    "学生：空手来（调查问卷当场发）。",
    "机房：只检查投影与音响；智多星本周不装。",
]:
    para(s, size=10.5, space_after=3)

# ================= 四、课堂流程 =================
h1("四、课堂流程（90 分钟 · 分钟级）")
make_table(
    ["时间", "环节", "教师做什么", "学生做什么", "设计意图"],
    [
        ("0–8", "生活悬念开场",
         "放「生活中的机器人」混剪；三问：①你家里哪个“自动”最聪明？②你被哪个“自动”坑过？③最想造个机器人帮你干什么？",
         "看视频；举手说", "从家里的事进门，不端不装"),
        ("8–18", "这门课学什么（大白话版）",
         "一句话：“学给机器人派活、定规矩”；怎么玩：2分/5分、同桌互审、随机抽人、记录册、彩蛋库；怎么考核：过程60%＋期末40%；强调“不背公式，靠想清楚”",
         "听＋提问", "规则透明，祛除对“编程课”的恐惧"),
        ("18–28", "发开学调查问卷",
         "发 03-E1 调查问卷（2 页）；讲清“不算分、摸底用”；当堂填、当堂交",
         "填问卷", "熟悉学生＋数学/逻辑摸底＋收集彩蛋素材"),
        ("28–33", "同桌互审定档＋抽人①",
         "宣布同桌互审规则：“同桌即搭档，一个做、一个查，做完交换”（方案 B 已定）；随机抽 2 人念第 3 题答案（最想让机器人干的活）",
         "定搭档；被抽者念答案", "定协作队形；熟悉抽人环节"),
        ("33–43", "机器人就是“照规矩干活”",
         "洗衣机＝选模式（如果…就…）＋转多久（做几遍）；扫地机＝没电回充（如果…就…）；售货机＝投币→选货→出货（先…再…）；结论：机器人聪明不聪明，全看规矩写得好不好",
         "跟着拆“洗衣机”一个例子", "用家里的事讲透课程内核"),
        ("43–48", "暂停点①（记录区）",
         "提问：“给扫地机器人加一条新规矩，让它更聪明”——先写在记录册第 1 页（记录册 05_ 在课件定稿后配套制作，"
         "与课件一一对应；印制前先写草稿纸）；不公布答案",
         "先写自己的猜测", "练记录册习惯：先猜后讲"),
        ("48–53", "抽人②",
         "抽 2–3 人念自己写的规矩；现场点评“哪条机器听得懂、哪条会出错”",
         "念＋听点评", "立规矩：写了就敢说"),
        ("53–68", "AI 智能体体验「今天来了一位 AI 同事」",
         "三步：①含糊命令“帮我整理这份名单”→结果乱；②清楚命令“按姓名排序、去掉重复、输出成表格”→结果漂亮；③起名建档（见名知意＋全班投票），生成“自我介绍＋岗位职责”",
         "观察两版对比；起名投票", "亲身体会：AI 听话，但要把话说清楚——和机器人一回事"),
        ("68–75", "学生当指挥",
         "同桌合作写一句“命令”让 AI 做件小事（写一条班级通知 / 整理一行乱数据）；抽 2 组上来现场试",
         "同桌讨论写命令；2 组演示", "建立信心：“我也能指挥”"),
        ("75–80", "暂停点②＋抽人③",
         "记录册第 2 页：“刚才两组的命令，哪句更好？差在哪？”先写；再抽人分享",
         "先写判断再听", "复盘“把话说清楚”的方法"),
        ("80–85", "财务里也有这种活",
         "打工/兼职的重复活（抄单、贴票、回消息）→财务里也一样（开票、对账、报销）→“下学期就做帮财务干活的机器人”；短视频 1 段",
         "听＋联想", "自然引到课程方向，零术语"),
        ("85–90", "收尾＋抽人④",
         "三句话总结；抽 3 人回答“这学期学什么”；布置课后任务；预告下周装智多星",
         "说／记任务", "制造下周期待"),
    ],
    [1.4, 2.9, 6.4, 3.0, 3.3], font_size=9)
para("时间合计：8+10+10+5+10+5+5+15+7+5+5+5 = 90 分钟。", size=10, bold=True, space_before=4)

# ================= 五、关键概念的生活化讲法 =================
h1("五、关键概念的生活化讲法（不背理论）")
for t in [
    ("概念 1　机器人＝照规矩干活的员工",
     "家里的例子：洗衣机「如果选快洗，就只洗 15 分钟」；扫地机器人「如果没电，就自己回充」；"
     "自动售货机「先投币、再选货、最后出货」。",
     "机器人没什么神秘：它就是一台照着规矩干活的机器。规矩写得好，它就像个聪明员工；"
     "规矩写得死板，它就闹笑话。",
     "本课程＝学怎么把规矩写得又清楚又灵活。"),
    ("概念 2　三种最基本的“规矩”（正式术语后面再学）",
     "顺序：先…再…（售货机）；条件：如果…就…（洗衣机的模式选择）；循环：一直做／做几遍（洗衣机转 30 分钟）。",
     "这三种规矩组合起来，就能让机器人干很复杂的活——本课程后面学的变量、命令，都是它们的升级版。"),
    ("概念 3　变量＝机器人装东西的盒子",
     "手机相册＝装照片的盒子；话费余额＝装数字的盒子；通讯录＝装名字和电话的盒子。",
     "机器人干活也要用盒子装东西：记金额、记名字、记日期——正式术语叫「变量」，项目三专门玩它。"),
    ("概念 4　AI 智能体＝会想事的机器人",
     "像手机里能帮你订票、写周报的 AI 助手：你让它做什么，它就做什么。",
     "但它只做「你清楚说出来的事」——没说的它要么不做、要么乱做。本课程学“把话说清楚”的方法，"
     "对机器人、对 AI 都一样。"),
]:
    para(t[0], size=11, bold=True, space_before=6, space_after=2)
    for line in t[1:]:
        para(line, size=10.5, space_after=2)
para("说明：本课不讲深，这四个说法够用；正式术语（变量、条件、循环）在项目二、三逐个“翻译”回来。",
     size=10, color=GRAY, space_before=4)

# ================= 六、趣味任务卡 =================
h1("六、趣味任务卡（第一课用）")
make_table(
    ["任务卡", "规则", "提示 / 彩蛋"],
    [
        ("A1 给 AI 员工起名建档\n难度★",
         "全班提名 → 投票定名＋形象；名字要求「见名知意」（一眼知道它干什么，比如“小财”“票票”）；提议被采纳者得 5 分",
         "提示：先立命名标准再提名，防止跑偏。\n彩蛋：名字和形象上「数字员工墙」，一学期沿用"),
        ("A2 指挥初体验\n难度★★",
         "同桌合作写一句「命令」让 AI 做件小事；抽 2 组现场演示；参与即 2 分，命令最清楚的组 5 分",
         "提示：说清楚「干什么＋怎么算干好」。\n彩蛋：胜组获「首席指挥官」称号"),
    ],
    [3.2, 8.0, 5.8], font_size=9.5)

# ================= 七、开学调查问卷 =================
h1("七、开学调查问卷（成品已单独生成：03-E1）")
para("问卷成品《03-E1_开学调查问卷_学生填写版_v1.0》已生成，2 页可打印。题目结构参照基础会计课"
     "「学生入口.html」的 21 题点单表，并按本课目标新增数学/逻辑摸底题。它只做三件事：熟悉学生、"
     "摸底数学/逻辑基础、为视频彩蛋收集素材——与记录册（05_）完全分开，各司其职。四大板块如下：",
     size=10.5, space_after=6)
make_table(
    ["板块", "题目要点", "用途"],
    [
        ("0 基本信息", "姓名 / 学号 / 班级；编程与机器人经历", "建档、分层"),
        ("一 数学与逻辑摸底", "数学自评；找规律；反推逻辑；步骤排序；“或者”判断；变量直觉（猜“盒子”）", "摸底，不算分"),
        ("二 电脑与 AI 接触", "编程小游戏；智能设备；用过的 AI", "调整讲法深浅"),
        ("三 彩蛋素材", "喜欢的动画角色；外号；口头禅；小习惯；最怕被问；三个词形容自己；同桌三个词；"
         "形象意愿；画风偏好；奖励偏好；惩罚偏好", "视频彩蛋与奖惩设计"),
        ("四 对课程的期待", "最想让机器人干的活；想对老师说的话", "案例选题"),
    ],
    [2.6, 9.4, 5.0], font_size=9.5)
para("照片沿用基础会计课已收照片（同一批学生）；若班级不同，另收一张。当场收齐，课后统计；"
     "第 2 节课前贴出「数字员工墙」（名字＋形象＋岗位职责）。",
     size=10.5, space_before=4)

# ================= 八、课后任务 =================
h1("八、课后任务")
for s in [
    "必做：用 30 秒给爸妈讲「这学期我要学什么」（口头即可；下节课抽 3 人复述，＋2 分）。",
    "选做①：带一条 AI 新闻下周分享（1 条＋2 分）。",
    "选做②：用任意 AI 生成一段「自我介绍」带来，全班玩「猜猜哪段是 AI 写的」（＋2 分）——把课堂的辨别乐趣带回家。",
]:
    para(s, size=10.5, space_after=3)

# ================= 九、教学备忘 =================
h1("九、教学备忘（易错点与应急）")
for s in [
    "大模型无网 / 卡顿：用 2 段预录屏兜底；两样都失效 → 跳过体验环节，宣布「下周我们亲手把这位同事请进来」，流程向后顺延。",
    "时间超支：「学生当指挥」压到 1 组；「财务里也有这种活」缩到 2 分钟。",
    "抽人冷场：被抽者答不出，同桌补答，第一课不扣分、不批评——先让学生不怕被抽。",
    "记录册：《我的机器人实验记录》（编号 05_）在第一讲课件定稿后配套制作，暂停点与课件一一对应；"
    "印制完成前，本课暂停点先写在草稿纸上，正式开课换用记录册对应页。",
    "视频素材：优先官方 / 自制素材，避免版权问题。",
]:
    para(s, size=10.5, space_after=3)

# ================= 十、修订记录 =================
h1("十、修订记录")
make_table(
    ["版本", "日期", "修订内容"],
    REVISION_HISTORY,
    [1.8, 2.6, 12.6], font_size=9.5)

def prune(prefix):
    """只保留最新版：删除同前缀、同扩展名的旧文件。"""
    for f in os.listdir(ROOT):
        if f.startswith(prefix) and f.endswith(".docx") and f != os.path.basename(OUT):
            os.remove(os.path.join(ROOT, f))


prune("03_项目一_第1讲_开学第一课_规划文档_")
doc.save(OUT)
print("已生成:", OUT)
