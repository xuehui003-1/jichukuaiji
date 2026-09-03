#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《03-E1_开学调查问卷_学生填写版》Word 文档（可打印 A4，2 页）。
题目结构参照基础会计课「学生入口.html」的 21 题点单表，并按本课目标：
①熟悉学生 ②摸底数学/逻辑思维基础 ③为视频彩蛋收集素材。
与记录册（05_）完全分开，各司其职。
版本号机制同规划文档：改 VERSION / DATE / REVISION_HISTORY 后重跑，只保留最新版。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============ 每次修订只需要改这里 ============
VERSION = "v1.0"
DATE = "2026-09-03"
REVISION_HISTORY = [
    ("v1.0", "2026-09-03",
     "新建：参照基础会计课「学生入口.html」21 题结构；新增数学/逻辑摸底题（找规律/反推逻辑/"
     "步骤排序/或者判断/变量直觉）与彩蛋素材题（外号/口头禅/小习惯/最怕被问/三个词/同桌三个词/"
     "形象意愿/画风/奖惩偏好）；照片沿用基础会计课已收照片。"),
]
# =============================================

BASE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE)
OUT = os.path.join(ROOT, "03-E1_开学调查问卷_学生填写版_{}_{}.docx".format(VERSION, DATE.replace("-", "")))

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2)
sec.top_margin = sec.bottom_margin = Cm(2)

GRAY = (89, 89, 89)
BLUE = (31, 78, 121)


def set_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr()
    rfonts = run._element.rPr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(text="", size=10.5, bold=False, align=None, name="宋体",
         color=None, space_after=4, space_before=0, line=1.3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, name=name, size=size, bold=bold, color=color)
    return p


def h1(text):
    para(text, size=13, bold=True, name="黑体", color=BLUE, space_before=8, space_after=4)


def h2(text):
    para(text, size=11.5, bold=True, name="黑体", color=(0, 0, 0), space_before=8, space_after=4)


def blank(n=46, lines=1):
    for _ in range(lines):
        para("＿" * n, size=10.5, color=GRAY, space_after=2, line=1.6)


def q(text, opts=None, answer_line=False, size=11):
    para(text, size=size, bold=True, space_before=4, space_after=2)
    if opts:
        para(opts, size=11, space_after=2)
    if answer_line:
        blank(40)


def add_footer():
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("03-E1 开学调查问卷 · {} · {} ｜ 不算分、摸底用 · 你写什么，这学期课上就出现什么".format(VERSION, DATE))
    set_font(r, size=9, color=GRAY)


add_footer()

# ================= 标题 =================
para("《财务机器人应用与开发》开学调查问卷", size=19, bold=True, name="黑体",
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=4)
para("这不是考试，没有标准答案，不算分。学号一定要填对。看不懂的题按字面答。当堂填、当堂交。",
     size=10.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
para("班级：＿＿＿＿＿＿＿＿　　姓名：＿＿＿＿＿＿＿＿　　学号：＿＿＿＿＿＿＿＿", size=11, space_after=2)
para("我的互审搭档（同桌）签名：＿＿＿＿＿＿＿＿　　"
     "照片：基础会计课已交过的不用再交；新同学交一张（彩蛋/抽人用）", size=10.5, space_after=8)

# ================= 0 基本信息 =================
h2("0　基本信息")
q("0-1　你接触过编程或机器人吗？",
  "□ 学过编程　　□ 玩过 Scratch / 编程小游戏　　□ 搭过乐高机器人　　□ 完全没有")

# ================= 一、数学与逻辑摸底 =================
h2("一、数学与逻辑摸底（不算分，做不出来也没关系）")
q("1　你对数学的感觉：", "□ 还不错　　□ 一般　　□ 比较吃力　　□ 看到数字就头大")
q("2　找规律：1、1、2、3、5、8、□　——　方框里填几？", answer_line=True)
q("3　小明说：“如果下雨，我就带伞。”今天小明没带伞。能断定今天没下雨吗？",
  "□ 能　　□ 不能　　□ 不确定")
q("4　泡一碗方便面，正确的顺序是：",
  "□ A　拆袋→烧水→倒调料→倒水→盖盖子\n□ B　烧水→拆袋→倒调料→倒水→盖盖子\n□ C　拆袋→倒调料→烧水→倒水→盖盖子")
q("5　晚会报名：“会唱歌或者会跳舞”的人可以报。小丽只会唱歌，她能报名吗？",
  "□ 能　　□ 不能　　□ 不确定")
q("6　机器人要记住钱数、日期、人名这些信息。你猜它至少需要几个“盒子”来装？",
  "随便写：＿＿＿＿ 个　　一句话说说为什么：")
blank(40)

# ================= 二、电脑与 AI =================
h2("二、电脑、机器人与 AI 接触")
q("7　家里 / 身边的智能设备：",
  "□ 智能音箱　　□ 扫地机器人　　□ 指纹/刷脸支付　　□ 都没有")
q("8　你用过哪些 AI：",
  "□ 语音助手（小爱、小度）　　□ AI 聊天　　□ AI 写文案　　□ AI 做图　　□ 都没用过")

# ================= 三、彩蛋素材 =================
h2("三、为课堂漫画 / 视频彩蛋提供素材")
q("9　你喜欢的动画 / 游戏角色（写 1–3 个）：", answer_line=True)
q("10　你的外号是什么？谁起的？", answer_line=True)
q("11　你的口头禅：", answer_line=True)
q("12　你上课有什么小动作 / 小习惯：", answer_line=True)
q("13　你最怕上课被问到什么：", answer_line=True)
q("14　用三个词形容你自己：", answer_line=True)
q("15　转头问同桌，让他用三个词形容你（让他写）：", answer_line=True)
q("16　愿不愿意被画进课堂漫画 / 视频彩蛋：",
  "□ 愿意参加，期待我的卡通形象　　□ 愿意参加，但想先看看示例\n□ 有点害羞，先参加一次试试看　　□ 这次先不参加")
q("17　如果画你，你喜欢哪种画风：",
  "□ 保留我平时的样子和几个小特点　　□ 和我喜欢的动画人物一起出现　　□ 做成适合当头像/表情包的样子\n□ 画成有趣的上课场景　　□ 请老师根据我的回答自由发挥")
q("18　做得好，想要什么奖励（可多选）：",
  "□ 画进上课漫画　　□ 卡通形象升级　　□ 加平时分　　□ 免一次作业　　□ 全班表扬　　□ 小零食")
q("19　表现不好，接受什么惩罚（可多选）：",
  "□ 做成搞笑表情包　　□ 上台讲一道题　　□ 帮同学讲懂一道题　　□ 表演才艺 30 秒　　□ 罚做一道题")

# ================= 四、对课程的期待 =================
h2("四、对这门课的期待")
q("20　你最想让机器人帮你干的活（可多选）：",
  "□ 抄表格　　□ 开发票　　□ 对账　　□ 整理文件　　□ 回消息　　□ 其他：＿＿＿＿＿＿")
q("21　关于这门课，想对老师说的话（选填）：")
blank(40, lines=2)

# 修订记录（脚注式，不占主版面）
para("", space_after=2)
p = para("修订记录：{}　{}　{}".format(VERSION, DATE, REVISION_HISTORY[0][2]), size=8, color=GRAY)

def prune(prefix):
    """只保留最新版：删除同前缀、同扩展名的旧文件。"""
    for f in os.listdir(ROOT):
        if f.startswith(prefix) and f.endswith(".docx") and f != os.path.basename(OUT):
            os.remove(os.path.join(ROOT, f))


prune("03-E1_开学调查问卷_学生填写版_")
doc.save(OUT)
print("已生成:", OUT)
