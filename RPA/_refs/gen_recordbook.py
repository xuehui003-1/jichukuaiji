# -*- coding: utf-8 -*-
"""
生成《05_项目一_第1讲_开学第一课_学生记录册》Word 文档。
v1.4 表格化改版（学基础会计 90-A1 记录册设计，配套 04_课件 v1.7）：
  结构＝暂停点①–⑤，每个暂停点三段：①题目描述一行 → ②三行空表格（我的想法/思考过程/老师的讲解）→ 下一个暂停点；
  空行全部改为带边框的空表格（左列标签＋右列作答），不再用下划线空行，也不出现嵌套编号；
  全册 1 页（标题三行＋5 个暂停点＋下课前自查），每人打印 1 页；
  封面语沿用基础会计记录册哲学：先写你的想法再听老师讲，猜错不扣分，猜错的地方正是今天要学的。
版本号机制同规划文档：改 VERSION / DATE / REVISION_HISTORY 后重跑，旧版自动删除（只保留最新版）。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ROW_HEIGHT_RULE

# ============ 每次修订只需要改这里 ============
VERSION = "v1.9"
DATE = "2026-09-06"
REVISION_HISTORY = [
    ("v1.9", "2026-09-06",
     "新增追加页 B/C/D（第 3/4/5 次课，配套 10_/12_/14_ 课件 v1.0）："
     "B＝变量进机器人（①打招呼顺序②先猜后验③读报错④排错三行⑤生活费三步）；"
     "C＝教机器人干活（①弹出提示三要素②三种串法对号＋家例③排错三行④收尾三句＋私例⑤外卖计算器图纸）；"
     "D＝跑起来修得好（①图纸三行②排错四行③验收单）。第 1 页与追加页 A 不动，已印的无需重印。"),
    ("v1.8", "2026-09-06",
     "追加页①②④随 08_ 课件 v2.2（费曼式引导链）改版：①改为「议一议——写 2 个有名字、会变的数＋"
     "1 个从来不变的，同桌互讲：它装什么、什么时候变」（电量引入之后的讨论暂停点）；"
     "②改为「余额格：早上 50→妈妈转来 200→买奶茶花 30，现在格子里是多少？50 和 200 去哪了」；"
     "④「五个货装五个盒」改「五个货各归哪类」（盒子说法全部弃用）；③⑤不动；第 1 页不动，已印的无需重印。"),
    ("v1.7", "2026-09-06",
     "追加页随 08_ 课件 v2.1 改版：①比喻更换——「盒子装数字」等牵强说法全部弃用，改用「贴了标签的储物格」"
     "（快递柜格：编号=名字、包裹=值、放新件旧件没、报编号取件）；"
     "②追加页暂停点①改为「收银机器人算一单要记住哪些东西」（名字＋记什么），替换原「手机里的盒子」；"
     "③其余②先猜后验／③上户口起名／④五类分类／⑤取件码争论不变；第 1 页不动，已印的无需重印。"),
    ("v1.6", "2026-09-06",
     "追加页随 08_ 课件 v2.0（变量专场·不碰软件）重排：项目二①你手机里的盒子／②先猜后验"
     "（盒子里先放35.5再放36.5，最后说出来的是几）／③变量上户口（给饭卡、话费、班费起名）／"
     "④五个货装五个盒／⑤取件码A8-2012能当数算吗（判断+理由，下节课机房揭晓）；"
     "排错三行挪到第 2 讲（进机房跑带病流程时用）；第 1 页不动，已印的无需重印。"),
    ("v1.5", "2026-09-06",
     "追加第 2 次课页（项目二第 1 讲）：①第 1 页（开学第一课暂停点①–⑤）与已印的 v1.4 一字不动，"
     "已印的学生册无需重印；②册后追加「第 2 次课 · 项目二第 1 讲」页——项目二暂停点①–⑤"
     "（积木数／手机里的盒子／先猜后跑／五个货装五个盒／排错三行），三行空表格制式不变；"
     "③第 1 页暂停点⑤「哪句 AI 命令更好」课堂上按新答案讲（两句都干成了，差在谁拿主意、出错能不能查），"
     "配套 04_课件 v1.8 与 07_讲课规划 v1.1。"),
    ("v1.4", "2026-09-04",
     "表格化改版（学基础会计 90-A1 记录册设计）：①每个暂停点＝题目描述一行＋三行空表格"
     "（左列「我的想法/思考过程/老师的讲解」＋右列作答），空行全部改为带边框空表格，"
     "删除下划线空行与嵌套编号，暂停点①–⑤依次排开不再叠套；②全册 1 页＋底部「下课前自查」；"
     "③封面语沿用基础会计：先写你的想法再听老师讲，猜错不扣分，猜错的地方正是今天要学的。"),
    ("v1.3", "2026-09-04",
     "暂停点制改版：全册 1 页、无分页符、无页眉，只收录暂停点，每处三行（①想法→②思考过程→③老师的讲解）。"),
    ("v1.2", "2026-09-04",
     "精简至 6 页，只记暂停点与老师检查点；填写区 1–2 行＋老师的答案区。"),
    ("v1.1", "2026-09-04",
     "标签一一对应版：页面按课堂顺序重排，页眉标「记录①–⑩」，填写项双区。"),
    ("v1.0", "2026-09-04",
     "新建，配套 04_课件 v1.3（学生视角版）。"),
]

# ============ 颜色 ============
ORANGE = (233, 120, 53)
ORANGE_DARK = (194, 90, 31)
BLUE = (59, 110, 165)
GRAY = (89, 89, 89)
BLACK = (0, 0, 0)

# ============ 输出 ============
BASE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE)
OUT = os.path.join(ROOT, "05_项目一_第1讲_开学第一课_学生记录册_{}_{}.docx".format(VERSION, DATE.replace("-", "")))

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(1.8)
sec.top_margin = sec.bottom_margin = Cm(1.5)


def set_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr()
    rfonts = run._element.rPr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(text="", size=10.5, bold=False, align=None, name="宋体",
         color=None, space_after=2, space_before=0, line=1.25):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, name=name, size=size, bold=bold, color=color)
    return p


def pause_point(no, title, desc, labels=("我的想法", "思考过程", "老师的讲解")):
    """一个暂停点：标题行 → 题目描述一行 → 三行空表格（左列标签＋右列作答）。"""
    # 标题行
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("⏸ 暂停点" + no + "　" + title)
    set_font(r, size=11.5, bold=True, name="黑体", color=ORANGE_DARK)
    # 题目描述一行
    para(desc, size=9.5, color=GRAY, space_after=2)
    # 三行空表格
    t = doc.add_table(rows=3, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, lab in enumerate(labels):
        row = t.rows[i]
        row.height = Cm(0.9)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        c0 = row.cells[0]
        c0.width = Cm(2.6)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(lab)
        set_font(r0, size=9.5, bold=True, name="黑体", color=BLACK)
        c1 = row.cells[1]
        c1.width = Cm(14.8)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)


def add_footer():
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("05_ 学生记录册 {} · {} ｜ 看到课件「📒暂停点N」才动笔".format(VERSION, DATE))
    set_font(r, size=8.5, color=GRAY)


add_footer()

# ================= 页首（标题三行） =================
para("我 的 机 器 人 实 验 记 录", size=18, bold=True, name="黑体",
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=1)
para("《财务机器人应用与开发》· 第 1 讲　　　　班级：＿＿＿＿　　学号：＿＿＿＿＿＿　　"
     "姓名：＿＿＿＿　　互审搭档：＿＿＿＿", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
para("这不是作业本，是实验记录。课上看到「📒暂停点N」才动笔：先写你的想法，再听老师讲；"
     "猜错不扣分——猜错的地方，正是你今天要学的地方。",
     size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

# ================= 暂停点①–⑤ =================
pause_point("①", "什么时候会「出事」？",
            "正常流程：放衣服→放洗衣液→关舱门→选模式→按启动→进水→洗涤→漂洗→脱水→蜂鸣结束。"
            "哪一步会出问题？把你想到的意外情况都写下来（越多越好）。")

pause_point("②", "洗衣机出事了①：停电",
            "正常：进水→洗涤→漂洗→脱水；万一：洗到一半突然停电。来电后，从哪一步继续？写下你的规矩。")

pause_point("③", "洗衣机出事了②：忘了放洗衣液",
            "正常：放洗衣液→进水洗涤；万一：忘了放洗衣液。机器怎么发现？怎么办？写下你的规矩。")

pause_point("④", "洗衣机出事了③：一直没人拿",
            "正常：洗完→蜂鸣提醒→取衣；万一：一直没人来拿。怎么办？写下你的规矩。")

pause_point("⑤", "哪句 AI 命令更好？差在哪？",
            "刚才两组上台给 AI 下命令。哪句命令更好？差在哪？先写判断，再听分享。")

# ================= 下课前自查 =================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("下课前自查：")
set_font(r, size=10, bold=True, name="黑体", color=ORANGE_DARK)
r2 = p.add_run("□ 暂停点①我写出了意外　□ ②③④我写了规矩　□ ⑤我写了哪句更好　　"
               "我猜错最惨的：＿＿＿＿＿＿＿＿＿＿")
set_font(r2, size=10, name="宋体")

# ================= 追加页：第 2 次课 · 项目二第 1 讲（v1.5 新增） =================
doc.add_page_break()
para("第 2 次课 · 项目二第 1 讲　变量——机器人的储物格　　　　日期：＿＿＿＿＿＿　　互审搭档：＿＿＿＿",
     size=12, bold=True, name="黑体", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("开场先补第 1 页的暂停点⑤（答案有新讲法），然后本页①–⑤随课上填。看到课件「📒暂停点N」才动笔。",
     size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

pause_point("①", "议一议：找 3 个「有名字的数」",
            "写 2 个「有名字、里面的数会变」的（例：游戏血条、钱包余额）；再写 1 个「有名字但从来不变」的（例：学号）。同桌互讲：它装什么、什么时候变。")

pause_point("②", "先猜后验：余额格早上 50，妈妈转来 200，买奶茶花 30——现在格子里是多少？50 和 200 去哪了？",
            "把你的猜测和理由写下来，再看黑板验证。猜错不扣分——猜错的地方正是今天要学的。",
            labels=("我的猜想", "验证之后", "老师的讲解"))

pause_point("③", "变量上户口：给饭卡余额、每月话费、班费各起一个变量名",
            "要求见名知义。写完同桌互审，不合格打回重起，合格的抄在下面。")

pause_point("④", "五个货，各归哪类？",
            "¥35.5／取件码 A8-2012／学生证.jpg／2026-09-09／14:30——分别属于：数值、字符、文件、日期、时间。先自己分，再听对答案。")

pause_point("⑤", "今天的争论题：取件码 A8-2012 能当数算吗？",
            "写下你的判断＋一条理由。今天不揭晓——下节课进机房跑一遍，把机器人的答案抄回来。",
            labels=("我的判断", "我的理由", "下节课的答案"))

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("下课前自查：")
set_font(r, size=10, bold=True, name="黑体", color=ORANGE_DARK)
r2 = p.add_run("□ ①我写了两个会变的数　□ ②我先猜后验　□ ③我起了三个名字　□ ④我分对了三种以上　"
               "□ ⑤我写了判断和理由　　我猜错最惨的：＿＿＿＿＿＿＿＿＿＿")
set_font(r2, size=10, name="宋体")

# ================= 追加页 B：第 3 次课 · 项目二第 2 讲（v1.9 新增） =================
doc.add_page_break()
para("第 3 次课 · 项目二第 2 讲　变量进机器人　　　　日期：＿＿＿＿＿＿　　互审搭档：＿＿＿＿",
     size=12, bold=True, name="黑体", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("机房课。看到课件「📒暂停点N」才动笔；翻车就写排错三行。",
     size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

pause_point("①", "《打招呼》里，机器人依次干了哪几件事？",
            "按顺序写（①②③…），一件一行，用你自己的话。写完和同桌对：排的顺序一样吗？")

pause_point("②", "先猜后验：放 50→输出，再放 200→输出——屏幕先后打出什么？50 和 200 都印吗？",
            "先写猜想再点运行。猜错不扣分——猜错的地方正是要点。",
            labels=("我的猜想", "验证之后", "老师的讲解"))

pause_point("③", "读报错：A8-2012 + 10 的报错原话",
            "把机器人说的原话一字不改抄下来，圈出读懂的词：它提到哪一行？什么类型？它想要什么？",
            labels=("机器人原话", "我读懂的词", "它想要的"))

pause_point("④", "排错三行（第一次用）：把翻车那次写下来",
            "这三行以后每次报错都写——机房铁律：翻车就写，写满一页换 2 分。",
            labels=("我在干什么", "机器人说了啥", "我改了什么·结果"))

pause_point("⑤", "生活费：月初 1500，吃饭花 600——纸上写三步",
            "①建什么变量（名字＋类型）②放什么 ③怎么算出剩多少。写完心里跑一遍——下节课开场要用！",
            labels=("变量清单", "放什么", "怎么算"))

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("下课前自查：")
set_font(r, size=10, bold=True, name="黑体", color=ORANGE_DARK)
r2 = p.add_run("□ ①我按顺序写全　□ ②我先猜后验　□ ③我抄了报错原话　□ ④我写满排错三行　□ ⑤我写了三步　　"
               "我最怕的报错词：＿＿＿＿＿＿＿＿＿＿")
set_font(r2, size=10, name="宋体")

# ================= 追加页 C：第 4 次课 · 项目二第 3 讲（v1.9 新增） =================
doc.add_page_break()
para("第 4 次课 · 项目二第 3 讲　教机器人干活——命令　　　　日期：＿＿＿＿＿＿　　互审搭档：＿＿＿＿",
     size=12, bold=True, name="黑体", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("机房课。规矩照旧：报错先自己读 3 秒；翻车就写排错三行。",
     size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

pause_point("①", "「弹出提示」的三要素",
            "动作＝干什么；参数＝按什么要求；输出＝出什么结果。空参数那次算不算有输出？写一句理由。",
            labels=("动作", "参数", "输出"))

pause_point("②", "三种串法对号入座＋你家的例子",
            "洗衣机（洗→漂→甩干）、地铁（超 10 公里 5 元、否则 3 元）、跑圈（满 3 圈才停）各是哪种串法？再各写一个你家里的例子。",
            labels=("顺序", "条件", "循环"))

pause_point("③", "排错三行：三条命令里你翻车的那次",
            "没翻车就写「差点翻车」的那次——三行是机房铁律。",
            labels=("我在干什么", "机器人说了啥", "我改了什么·结果"))

pause_point("④", "收尾三句＋你的私例",
            "①命令＝动作＋参数＋输出　②三种串法：顺序/条件/循环　③参数填不对，命令就歪——每句后面配一个你自己的例子。")

pause_point("⑤", "期末图纸：外卖满减计算器（单价 39.9、买 2 份、满 25 减 4）",
            "①要哪几个变量（名字＋类型）②哪几条命令 ③用哪种串法。下节课照图纸开工——图纸不全不开机。",
            labels=("变量清单", "命令清单", "串法"))

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("下课前自查：")
set_font(r, size=10, bold=True, name="黑体", color=ORANGE_DARK)
r2 = p.add_run("□ ①我会拆三要素　□ ②我对上三种串法　□ ③我写了排错三行　□ ④三句配了私例　□ ⑤图纸画全了　　"
               "家里例子最妙的一个：＿＿＿＿＿＿＿＿＿＿")
set_font(r2, size=10, name="宋体")

# ================= 追加页 D：第 5 次课 · 项目二第 4 讲（v1.9 新增） =================
doc.add_page_break()
para("第 5 次课 · 项目二第 4 讲　跑起来修得好——综合与调试　　　　日期：＿＿＿＿＿＿　　互审搭档：＿＿＿＿",
     size=12, bold=True, name="黑体", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("项目二收官。没图纸不开机；卡住三板斧：读报错→找组件→问同桌。",
     size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

pause_point("①", "图纸三行（补全＋同桌点头才许开机）",
            "①变量清单（名字＋类型）②命令清单 ③串法。讲给同桌听一遍，他点头才开机。",
            labels=("变量清单", "命令清单", "串法"))

pause_point("②", "排错四行（升级版）：今天真翻的那次",
            "比机房第一天多一行：改完之后怎样了——这是你的收网记录。",
            labels=("现象·机器人说啥", "我改了什么", "改完之后"))

pause_point("③", "主管验收单",
            "他的流程跑给你看，再出个刁钻输入考它：0 份？恰好 25 元？负数？翻车的帮他指到那一步。",
            labels=("我测了什么", "它的表现", "我建议他改哪"))

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("下课前自查：")
set_font(r, size=10, bold=True, name="黑体", color=ORANGE_DARK)
r2 = p.add_run("□ ①图纸点头了　□ ②我写全四行　□ ③我出了刁钻输入　　"
               "今天修掉的最狠的 bug：＿＿＿＿＿＿＿＿＿＿")
set_font(r2, size=10, name="宋体")


def prune(prefix):
    """只保留最新版：删除同前缀、同扩展名的旧文件。"""
    for f in os.listdir(ROOT):
        if f.startswith(prefix) and f.endswith(".docx") and f != os.path.basename(OUT):
            os.remove(os.path.join(ROOT, f))


prune("05_项目一_第1讲_开学第一课_学生记录册_")
doc.save(OUT)
print("已生成:", OUT)
