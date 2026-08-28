#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《02_项目四_自动化处理_讲课文档_试样》Word 文档。
展示讲课文档的八段式固定结构与分钟级课堂流程；附变量/命令/Excel/浏览器/邮件/OCR 趣味任务样例库。
版本号机制同规划文档：改 VERSION / DATE / REVISION_HISTORY 后重跑，旧版保留。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============ 每次修订只需要改这里 ============
VERSION = "v1.0"
DATE = "2026-08-26"
REVISION_HISTORY = [
    ("v1.0", "2026-08-26",
     "新建试样：展示讲课文档八段式固定结构；第1讲 Excel 自动化完整试样；"
     "浏览器/邮件/屏幕录制/OCR 要点试样；附跨项目趣味任务样例库 13 项。"),
]
# =============================================

BASE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE)
OUT = os.path.join(ROOT, "02_项目四_自动化处理_讲课文档_试样_{}_{}.docx".format(VERSION, DATE.replace("-", "")))

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2)
sec.top_margin = sec.bottom_margin = Cm(2)

GRAY = (89, 89, 89)
BLUE = (31, 78, 121)


def set_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr()
    rfonts = run._element.rPr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(text="", size=10.5, bold=False, align=None, name="宋体",
         color=None, space_after=4, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.3
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, name=name, size=size, bold=bold, color=color)
    return p


def h1(text):
    para(text, size=15, bold=True, name="黑体", color=BLUE, space_before=12, space_after=6)


def h2(text):
    para(text, size=12, bold=True, name="黑体", color=(0, 0, 0), space_before=8, space_after=4)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def make_table(headers, rows, widths_cm, font_size=9.5, header_bg="D9E2F3"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(htxt)
        set_font(r, size=font_size, bold=True)
        shade_cell(hdr[i], header_bg)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(v))
            set_font(r, size=font_size)
    for i, w in enumerate(widths_cm):
        for row in t.rows:
            row.cells[i].width = Cm(w)
    return t


def add_footer():
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("第 ")
    set_font(r1, size=9, color=GRAY)
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(it); run._r.append(f2)
    set_font(run, size=9, color=GRAY)
    r2 = p.add_run(" 页 · 共 ")
    set_font(r2, size=9, color=GRAY)
    run2 = p.add_run()
    g1 = OxmlElement("w:fldChar"); g1.set(qn("w:fldCharType"), "begin")
    it2 = OxmlElement("w:instrText"); it2.set(qn("xml:space"), "preserve"); it2.text = "NUMPAGES"
    g2 = OxmlElement("w:fldChar"); g2.set(qn("w:fldCharType"), "end")
    run2._r.append(g1); run2._r.append(it2); run2._r.append(g2)
    set_font(run2, size=9, color=GRAY)
    r3 = p.add_run(" 页")
    set_font(r3, size=9, color=GRAY)


add_footer()

# ================= 标题区 =================
para("项目四　自动化处理（讲课文档 · 试样）", size=20, bold=True, name="黑体",
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, space_before=6)
para("财务机器人应用与开发 · 第 7–9 周 · 6 课时", size=12, name="黑体",
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("文档编号 02 ｜ 版本 {} ｜ 生成日期 {} ｜ 配套《00_财务机器人课程教学规划_v3.1》｜ 性质：格式试样".format(VERSION, DATE),
     size=10.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# ================= 0 试样说明 =================
h1("0　这份试样看什么（请先读）")
for s in [
    "① 本文件是讲课文档的「格式试样」：只把第 1 讲 Excel 自动化完整写出来（第三部分），"
    "其余知识点给要点——您先看格式和深度对不对，再决定正式版怎么写。",
    "② 四个重点请反馈：固定结构喜不喜欢、课堂流程细不细、趣味任务卡长这样行不行、"
    "整体风格是不是您要的「有趣但不飘」。",
    "③ 角色说明：试样暂用「李会计、王老板」两个人物（与您基础会计课衔接），正式版按角色机制文档另定。",
    "④ 正式版每份讲课文档会配套：讲课文档（主文件）+ 学生任务单（可打印）+ 课堂用数据文件包。",
]:
    para(s, size=10.5, space_after=3)

# ================= 一、总览 =================
h1("一、项目四总览")
make_table(
    ["知识点", "课时", "智多星主要组件", "趣味任务锚点"],
    [
        ("Excel 自动化", "2", "打开/读取单元格、写入、循环、条件判断", "E1 工资条生成器、E2 报销表打假、E3 Excel 找不同"),
        ("浏览器操作自动化", "1", "打开网页、元素定位、点击、输入、提取", "E4 汇率小秘书、E5 查分机器人"),
        ("邮件处理自动化", "1", "收发邮件、读附件、筛选转发", "E6 工资条小邮差、E7 报销分拣员"),
        ("屏幕录制", "1", "录制、回放、导出", "E8 给学弟学妹的 SOP 视频"),
        ("OCR 识别", "1", "OCR 组件、结果校验", "E9 发票照妖镜、E10 识别大赛"),
    ],
    [3.0, 1.0, 5.0, 8.0], font_size=10)

# ================= 二、固定结构 =================
h1("二、讲课文档的固定结构（正式版每份都按这个来）")
para("0 使用说明（这份怎么用）→ 1 本讲目标 → 2 课前准备 → 3 课堂流程（分钟级）→ "
     "4 知识点详解（是什么·为什么·怎么做·常见坑）→ 5 趣味任务卡 → 6 课后任务 → "
     "7 教学备忘（易错点与应急）。", size=11, bold=True, space_after=3)
para("试样中「第 1 讲」完整走完这 8 段；其余知识点只给第 1、5、7 段要点，省篇幅。", size=10.5, space_after=3)

# ================= 三、第1讲 Excel 自动化 =================
h1("三、第 1 讲 Excel 自动化（完整试样 · 90 分钟）")

h2("3.1　本讲目标")
for s in [
    "知识：理解 Excel 自动化解决什么问题；知道「区域、单元格、循环、条件判断」四个核心概念。",
    "技能：会搭一个「读取 → 处理 → 写回」流程；能改造成工资条生成器。",
    "彩蛋：任务完成者解锁「数字员工·小E」贴纸一枚。",
]:
    para(s, size=10.5, space_after=3)

h2("3.2　课前准备")
for s in [
    "机房：智多星离线版 + Excel；教师机大屏。",
    "数据文件：①《报销明细表_乱版.xlsx》（100 行，藏 3 个错：金额混入文字、合计不对、有空行）"
    "②《工资表_一行版.xlsx》③空白《汇总表.xlsx》。",
    "分组：4 人一组——组长（操作）、副操作（复核）、记录员（写暂停点答案）、播报员（展示讲解）。",
]:
    para(s, size=10.5, space_after=3)

h2("3.3　课堂流程（90 分钟）")
make_table(
    ["时间", "环节", "教师做什么", "学生做什么", "设计意图"],
    [
        ("0–5", "引入", "放「李会计的周一噩梦」漫画：100 张报销单要汇总，加班到九点", "看图回答「他为什么累」", "用场景立靶子"),
        ("5–20", "演示", "当堂搭「读取单元格→求和→写回汇总表」，边搭边讲", "观察并把步骤记到记录单", "先见完整成品"),
        ("20–23", "暂停点1", "问：机器人处理 100 行和 1000 行，时间差多少？", "先写猜测，再听讲解", "预测效应：先猜后讲记得牢"),
        ("23–38", "讲解", "循环结构（按行遍历）、条件判断（超标标红）", "对照自己的猜测找差距", "补全「循环」核心概念"),
        ("38–53", "任务1", "发《工资表_一行版》，要求改造成工资条（加插入行）", "小组协作，跑通为准", "从模仿到改造"),
        ("53–55", "暂停点2", "问：机器人为什么不会累、不会看串行？", "小组讨论，组长发言", "引出机器人的价值点"),
        ("55–70", "任务2", "《报销明细表_乱版》打假：自动找错行、标红、重算合计", "先人肉找错 1 分钟，再让机器人找，对比", "体验人机差距"),
        ("70–80", "展示投票", "组织投票：最快跑通组、最稳流程组", "播报员演示，其他组挑刺", "展示即考核"),
        ("80–90", "收尾", "小结三个坑（区域选择、格式陷阱、死循环）", "记入记录单；领课后任务", "防课后翻车"),
    ],
    [1.3, 1.7, 5.0, 5.0, 3.2], font_size=9)

h2("3.4　知识点详解（是什么 · 为什么 · 怎么做 · 常见坑）")
for t in [
    ("知识点 1　打开与读取 Excel",
     "是什么：让机器人打开工作簿、读指定单元格或区域。",
     "为什么：Excel 是财务人最主要的「桌面」，机器人先得会读它。",
     "怎么做：选文件路径 → 选工作表 → 读单元格（行号列号或 A1 写法）。",
     "常见坑：文件被人工打开时机器人读不到（占用）；中文路径没问题，但少一个反斜杠就报错。"),
    ("知识点 2　写入与保存",
     "是什么：把结果写回单元格、保存为新文件。",
     "为什么：机器人处理完必须留痕；原表不能覆盖。",
     "怎么做：结果写进新文件，原文件只读。",
     "常见坑：覆盖原表后无法找回；忘加「保存」组件，流程跑完等于没跑。"),
    ("知识点 3　循环（按行遍历）",
     "是什么：对每一行重复执行同一组命令。",
     "为什么：100 行和 10000 行对机器人是同一件事。",
     "怎么做：取总行数 → 计数器从 1 到 N → 循环体里用「当前行」读写。",
     "常见坑：循环条件写反 → 死循环；循环体里忘了「行号 +1」。"),
    ("知识点 4　条件判断（如果…否则…）",
     "是什么：满足条件走 A 分支，否则走 B。",
     "为什么：标红超标、筛出异常全靠它。",
     "怎么做：如果 金额>1000 标红 否则 不动。",
     "常见坑：条件写反把正常行标红；大于号与等号混用。"),
]:
    para(t[0], size=11, bold=True, space_before=6, space_after=2)
    for line in t[1:]:
        para(line, size=10.5, space_after=2)

h2("3.5　趣味任务卡（第 1 讲用）")
make_table(
    ["任务卡", "规则", "提示 / 彩蛋"],
    [
        ("E1《工资条生成器》\n难度★★",
         "把「一行版工资表」变成每人带表头的工资条；跑通且格式正确得 5 分",
         "提示：插入行组件放进循环里。\n彩蛋：第一名小组下节课获「先手券」（任一任务提前 2 分钟开跑）"),
        ("E2《报销表打假》\n难度★★★",
         "机器人找出 3 处错（金额混文字、合计错、空行），标红并重算合计；找全得 5 分",
         "提示：先人肉找 1 分钟，再用机器人找。\n彩蛋：找全组获「审计天团」称号"),
        ("E3《Excel 找不同》\n（课后）难度★★★",
         "AI 生成两版销售表，机器人自动比对差异，出一份《差异报告》，下节课前交",
         "提示：逐列对比数值。\n彩蛋：报告最规范者下周当「课代表」"),
    ],
    [3.2, 8.0, 5.8], font_size=9.5)

h2("3.6　课后任务")
para("必做：E3 Excel 找不同。选做：用 3 句话给爸妈讲明白「什么是循环」。", size=10.5, space_after=3)

h2("3.7　教学备忘")
para("易错点：死循环（先教学生按「停止」，再讲原因）；区域选错（记住 A1 记法）。", size=10.5, space_after=2)
para("应急：有小组跑不通 → 教师机投屏集体排障一次——这就是最好的现场教学。", size=10.5, space_after=2)

# ================= 四、第2讲 浏览器 =================
h1("四、第 2 讲 浏览器操作自动化（要点）")
h2("4.1　目标")
para("会用「打开网页、元素定位、点击、输入、提取」五类组件。", size=10.5, space_after=3)
h2("4.2　课堂流程要点")
para("引入「机器人帮我查快递」→ 演示自动抓汇率 → 任务：做一台汇率小秘书 → 收尾讲两个坑"
     "（页面没加载完就点击、验证码——预告项目六的验证码识别）。", size=10.5, space_after=3)
h2("4.3　趣味任务卡")
make_table(
    ["任务卡", "规则", "提示 / 彩蛋"],
    [
        ("E4《汇率小秘书》\n难度★★",
         "每天自动抓当日汇率写入 Excel 并画折线图",
         "彩蛋：数据最全组获「外汇交易员」称号"),
        ("E5《查分机器人》\n难度★★★",
         "自动登录教务系统查成绩、截图汇总",
         "合规提醒：只查自己的账号；讲清「等待元素出现」\n彩蛋：全班票选「最想要的机器人」"),
    ],
    [3.2, 8.0, 5.8], font_size=9.5)
para("衔接：企业信息查验是项目六的主任务，这里先学会「打开、输入、提取」三件套。",
     size=10.5, space_before=3)

# ================= 五、第3讲 邮件 =================
h1("五、第 3 讲 邮件处理自动化（要点）")
h2("5.1　目标")
para("会用「收发邮件、读取附件、筛选转发」组件。", size=10.5, space_after=3)
h2("5.2　课堂流程要点")
para("引入「工资条发错人」的新闻事故 → 演示自动发邮件 → 任务：工资条小邮差 → "
     "收尾讲两个坑（附件路径写死、收件人列表错位）。", size=10.5, space_after=3)
h2("5.3　趣味任务卡")
make_table(
    ["任务卡", "规则", "提示 / 彩蛋"],
    [
        ("E6《工资条小邮差》\n难度★★★",
         "Excel 每人一行 → 机器人逐个生成邮件发出；错发一人扣 2 分",
         "发给「模拟同事」邮箱列表。\n彩蛋：零事故组获「金牌小邮差」"),
        ("E7《报销分拣员》\n难度★★",
         "收到带附件的报销邮件 → 自动下载附件、按日期重命名、归类文件夹、自动回复「已收到」",
         "彩蛋：写得最有礼貌的自动回复上墙展示"),
    ],
    [3.2, 8.0, 5.8], font_size=9.5)

# ================= 六、第4-5讲 屏幕录制与OCR =================
h1("六、第 4–5 讲 屏幕录制与 OCR（要点）")
for s in [
    "屏幕录制：用途＝把操作录成 SOP 手册。任务卡 E8《给学弟学妹的 SOP 视频》★★——录自己的流程操作，"
    "剪映加字幕配音，优秀作品进班级彩蛋库（衔接说课任务库 T1/T2 配音玩法）。",
    "OCR：用途＝识别发票、单据。任务卡 E9《发票照妖镜》★★——拍照→识别→自动填 Excel，识别错处追因"
    "（模糊/倾斜/遮挡）；E10《识别大赛》★★★——小组 PK 准确率与速度，输的组才艺表演。",
    "衔接：OCR 输出交给 AI 审核＝本项目的「三个数字员工一台戏」点睛活动。",
]:
    para(s, size=10.5, space_after=3)

# ================= 七、任务样例库 =================
h1("七、附：趣味任务样例库（变量 / 命令 / Excel / 浏览器 / 邮件 / OCR）")
para("以下任务横跨项目三、项目四，正式版会拆到各项目讲课文档的任务卡里；这里集中列出供您先挑。",
     size=10.5, space_after=6)
make_table(
    ["编号", "主题", "任务名", "一句话玩法", "难度", "用在哪"],
    [
        ("F1", "变量", "变量上户口", "新建变量＝给机器人发身份证：名字必须「见名知意」（auditAmount、voucherDate），起名审核制，审核不过打回", "★", "项目三"),
        ("F2", "变量", "变量分类大乱斗", "PPT 快速闪现 10 个财务数据（¥35,600.00、2026-08-26、「办公室租金」、D:\\凭证\\8月.xlsx、14:30…），学生举牌分类", "★", "项目三"),
        ("F3", "命令", "命令抽卡搭流程", "智多星常用命令印成卡牌，小组抽卡限时拼出指定流程，拼错现场报错演示", "★★", "项目三"),
        ("F4", "命令", "报错医生", "3 段带病流程（金额写成字符型、循环条件写反、路径少反斜杠），学生写「诊断书」", "★★★", "项目三"),
        ("F5", "Excel", "工资条生成器", "循环 + 插入行，把一行工资表变成每人带表头的工资条", "★★", "项目四"),
        ("F6", "Excel", "报销表打假", "机器人找错行、标红、重算合计，人机 PK", "★★★", "项目四"),
        ("F7", "Excel", "Excel 找不同", "AI 生成两版销售表，机器人自动比对出差异报告", "★★★", "项目四"),
        ("F8", "浏览器", "汇率小秘书", "每天自动抓汇率填入 Excel 画折线图", "★★", "项目四"),
        ("F9", "浏览器", "查分机器人", "自动登录教务系统查成绩截图汇总（合规：只查自己账号）", "★★★", "项目四"),
        ("F10", "邮件", "工资条小邮差", "Excel 每人一行逐个发邮件，错发一人扣 2 分", "★★★", "项目四"),
        ("F11", "邮件", "报销分拣员", "自动下载附件、按日期重命名归类、自动回复", "★★", "项目四"),
        ("F12", "OCR", "发票照妖镜", "拍照识别自动填 Excel，识别错处追因", "★★", "项目四"),
        ("F13", "综合", "机器人新品发布会", "期末路演：机器人演示 + 答辩 + 全班投票（见规划文档任务库 T3）", "★★★", "结课"),
    ],
    [1.0, 1.4, 2.6, 8.0, 1.1, 2.0], font_size=9)

# ================= 八、反馈清单 =================
h1("八、给老师的确认清单（试样反馈 4 问）")
for s in [
    "□ 固定结构（0–7 八段）是否合适？",
    "□ 第 1 讲写到的「分钟级流程 + 暂停点」是您要的颗粒度吗？",
    "□ 趣味任务卡的形式（规则 / 提示 / 彩蛋 / 评分）喜不喜欢？",
    "□ 风格是否「有趣但不飘」？哪些地方要更活泼、哪些要更严肃？",
]:
    para(s, size=10.5, space_after=3)
para("反馈后助手按意见出正式版：先出项目三，再依次出项目四至七，每份独立编号（03_/04_/…）并自动带版本号。",
     size=10.5, bold=True, space_before=4)

# ================= 修订记录 =================
h1("九、修订记录")
make_table(
    ["版本", "日期", "修订内容"],
    REVISION_HISTORY,
    [1.8, 2.6, 12.6], font_size=9.5)

doc.save(OUT)
print("已生成:", OUT)
