#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《基础会计》学生记录册 Word 版（可编辑）
- 关键暂停点：三格（暂停点答案 / 我是这么想的 / 老师讲完之后）
- 普通暂停点：一格（暂停点答案）+ 一行「老师讲完之后」窄格
- 首页并入《第一课 算不清的账》精简版
用法：python3 build_记录册_docx.py
"""
import json, re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SRC  = os.path.join(HERE, "90-A1_学生记录册_全学期_可打印_v15.0_20260830.html")
OUT  = os.path.join(HERE, "90-A1_学生记录册_全学期_可打印_v17.0_20260901.docx")

# ---------- 从 v15 HTML 抽取暂停点 ----------
import html as htmllib
def extract(path):
    s = open(path, encoding="utf-8").read()
    def clean(x):
        x = re.sub(r"<br\s*/?>", "\n", x or "")
        x = re.sub(r"<[^>]+>", "", x)
        return htmllib.unescape(x).strip()
    data, cur = [], None
    for m in re.finditer(r"<h2>(.*?)</h2>|<article class=\"pause\">(.*?)</article>", s, re.S):
        if m.group(1):
            cur = {"proj": clean(m.group(1)), "items": []}
            data.append(cur)
        else:
            blk = m.group(2)
            h3 = re.search(r"<h3>(.*?)</h3>", blk, re.S)
            q  = re.search(r'<p class="q">(.*?)</p>', blk, re.S)
            if cur is None:
                cur = {"proj": "未分组", "items": []}; data.append(cur)
            cur["items"].append({
                "t": clean(h3.group(1)) if h3 else "",
                "q": clean(q.group(1)) if q else "",
                "table": bool(re.search(r"<table", blk)),
            })
    return data

# 关键暂停点判定：需要解释、推理、用自己的话 → 值得写三格
KEY   = re.compile(r"★加分题|为什么|用自己的话|你觉得|能不能归成|如果你是|写一条理由|怎么办|你会怎么|猜|还能按什么|到底是什么意思")
LIGHT = re.compile(r"抄下来|抄在这里|连线|排序|分别叫什么|把你填的")
def is_key(it):
    t = it["t"] + " " + it["q"]
    return bool(KEY.search(t)) and not bool(LIGHT.search(t))

# ---------- Word 样式工具 ----------
FONT = "微软雅黑"
def setfont(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def set_borders(tbl, color="BFBFBF", sz=6):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"),"single"); e.set(qn("w:sz"),str(sz))
        e.set(qn("w:space"),"0");    e.set(qn("w:color"),color)
        borders.append(e)
    tblPr.append(borders)

def writing_box(doc, label, lines=3, fill="FFFFFF"):
    """一个书写框：标签 + 带下划线的空白区"""
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    setfont(p.add_run(label), 9.5, bold=True, color=(0x44,0x44,0x44))
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t)
    c = t.cell(0,0); shade(c, fill)
    c.text = ""
    for i in range(lines):
        pp = c.paragraphs[0] if i==0 else c.add_paragraph()
        pp.paragraph_format.space_after = Pt(0)
        pp.paragraph_format.line_spacing = Pt(19)
        setfont(pp.add_run(""), 10.5)
    return t

# ---------- 开始生成 ----------
doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.6); sec.bottom_margin = Cm(1.6)
sec.left_margin = Cm(1.7); sec.right_margin = Cm(1.7)

st = doc.styles["Normal"]
st.font.name = FONT; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def H(text, size=16, color=(0x1a,0x23,0x7e), before=10, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    setfont(p.add_run(text), size, bold=True, color=color)
    return p

def P(text, size=10.5, color=(0x33,0x33,0x33), bold=False, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    setfont(p.add_run(text), size, bold=bold, color=color)
    return p

# ===== 封面 =====
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
setfont(p.add_run("我的会计实验记录"), 30, bold=True, color=(0x1a,0x23,0x7e))
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
setfont(p.add_run("《基础会计》全学期 · 这不是作业本，是实验记录"), 12, color=(0x66,0x66,0x66))

t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
set_borders(t)
for i, lab in enumerate(["姓名", "学号", "班级"]):
    c = t.cell(0,i); c.text = ""
    pp = c.paragraphs[0]; pp.paragraph_format.space_before = Pt(10)
    pp.paragraph_format.space_after = Pt(10)
    setfont(pp.add_run(lab + "：____________"), 11)

P("")
P("怎么用这本册子", 13, (0x1a,0x23,0x7e), bold=True, after=6)
P("课上会停下来问你。看到「暂停点」就翻到对应编号，把你想的写下来。编号跟投影上一样：项目几第几暂停点。", 10.5)
P("★ 关键暂停点写三格：暂停点答案 / 我是这么想的 / 老师讲完之后，把那个办法写下来。", 10.5)
P("○ 其余暂停点写一格就够：把答案写下来；老师讲完如果和你不一样，再补一行。", 10.5)
P("猜错不扣分。写了 2 分，写好 5 分。一道题最多 5 分。册子不用交，一直带着。", 10.5, (0x88,0x00,0x00))

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ===== 第一课：算不清的账（精简版并入） =====
H("第一课　算不清的账", 20)
P("四道题，全是小学算术——加减乘除，没有别的。不打开课本。先自己算，不许讨论。", 11, (0x66,0x66,0x66))
P("每年都有一大半人第一遍算错。算错很正常，这正是今天要讲的事。", 11, (0x88,0x00,0x00))

D1 = [
    ("第一题　校门口的小卖部",
     "顾客买一瓶饮料（进价 3 元，卖 5 元），掏出一张 100 元。老板没零钱，到隔壁把 100 元换成零钱，"
     "找给顾客 95 元。晚上隔壁找上门：那张 100 元是假的。老板赔了隔壁 100 元。\n"
     "问：小卖部老板这一天一共亏了多少？"),
    ("第二题　宿舍四个人去超市",
     "你垫了 180 元，小李垫了 60 元，小王垫了 30 元，小张没带钱一分没出。四个人说好平分。\n"
     "问：最后谁给谁多少钱？"),
    ("第三题　奶茶店充 100 送 50",
     "充 100 元，卡里变成 150 元。\n问：这相当于打几折？顺手再算：充 100 送 20 是几折？充 500 送 300 呢？"),
    ("第四题　校外兼职，两个坑",
     "坑一：说好时薪 15 元、每天 8 小时。实际每天要待 10 小时，老板说准备和收拾「不算工作时间」。\n"
     "问①：你的实际时薪是多少？\n"
     "坑二：另一个兼职日薪 100，必须干满 7 天才结账，中途走不给钱。你干了 3 天，不想干了。\n"
     "问②：走，还是不走？为什么？"),
]
for title, body in D1:
    H(title, 13, (0x0d,0x47,0xa1), before=12, after=4)
    P(body, 10.5)
    writing_box(doc, "暂停点答案：", 2)
    writing_box(doc, "我是这么想的：（写、画、列表都行）", 3, "FAFAFA")
    writing_box(doc, "老师讲完之后，把那个办法写下来：", 2)

H("课后　把这个办法用在自己身上一次", 13, (0x0d,0x47,0xa1), before=12)
P("找一件你算不清、或者算错过的钱：一次 AA、一次充值满减、一次兼职代购都行。按下面四格填，周日前发班级群。", 10.5)
P("四格都填＝合格。③里有算式＝5 分。只写感想不写数字＝退回重写。", 10.5, (0x88,0x00,0x00), bold=True)
writing_box(doc, "① 那件事（一句话说清）：", 2)
writing_box(doc, "② 我当时以为是多少：", 1, "FAFAFA")
writing_box(doc, "③ 用今天的办法重算（要写算式）：", 3)
writing_box(doc, "④ 差了多少：", 1, "FAFAFA")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
setfont(p.add_run("碰到钱的事，先换个位置看一眼。"), 13, bold=True, color=(0x88,0x00,0x00))
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
setfont(p.add_run("今天只记住这一句就够了"), 9.5, color=(0x99,0x99,0x99))

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ===== 全学期暂停点 =====
data = extract(SRC)
nkey = nlite = 0
for g in data:
    if not g["items"]:
        continue
    H(g["proj"], 18)
    for it in g["items"]:
        key = is_key(it)
        nkey += key; nlite += (not key)
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(2)
        mark = "★ " if key else "○ "
        setfont(p.add_run(mark + it["t"]), 12, bold=True,
                color=(0xc6,0x28,0x28) if key else (0x0d,0x47,0xa1))
        if it["q"]:
            P(it["q"], 10.5, (0x22,0x22,0x22), after=2)

        if it["table"]:   # 项目1第13暂停点：五行表
            tb = doc.add_table(rows=6, cols=5); set_borders(tb)
            hdr = ["第几天", "发生了什么", "左边怎么变", "右边怎么变", "歪吗"]
            rows = ["第1天 投入 ____ 万", "第3天 买奶茶机 ____ 万", "第5天 借款 ____ 万",
                    "第10天 进货 ____ 万 未付", "第20天 还款 ____ 万"]
            for j, h in enumerate(hdr):
                c = tb.cell(0,j); c.text=""; shade(c, "E8EAF6")
                setfont(c.paragraphs[0].add_run(h), 9.5, bold=True)
            for i, r in enumerate(rows, start=1):
                a, b = r.split(" ", 1)
                c = tb.cell(i,0); c.text=""; setfont(c.paragraphs[0].add_run(a), 9.5)
                c = tb.cell(i,1); c.text=""; setfont(c.paragraphs[0].add_run(b), 9.5)
                for j in (2,3,4):
                    cc = tb.cell(i,j); cc.text=""
                    cc.paragraphs[0].paragraph_format.line_spacing = Pt(18)
                    setfont(cc.paragraphs[0].add_run(""), 9.5)

        if key:
            writing_box(doc, "暂停点答案：", 2)
            writing_box(doc, "我是这么想的：", 2, "FAFAFA")
            writing_box(doc, "老师讲完之后，把那个办法写下来：", 2)
        else:
            writing_box(doc, "暂停点答案：", 2)
            writing_box(doc, "老师讲完之后（和你不一样再补）：", 1, "FAFAFA")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ===== 学期末 =====
H("学期末，回头看看", 18)
for q in ["这学期我最有把握的一件事：", "我还是搞不清的一件事：", "如果重来一次，我会："]:
    writing_box(doc, q, 3)

doc.save(OUT)
print(f"已生成：{os.path.basename(OUT)}")
print(f"  关键暂停点（三格）：{nkey}")
print(f"  普通暂停点（一格）：{nlite}")
print(f"  合计：{nkey+nlite}")
