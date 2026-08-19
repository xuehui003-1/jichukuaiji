# -*- coding: utf-8 -*-
"""主文档 vs 记录册 一致性自检 —— 以后每做完一个项目都跑这个"""
import os, re, glob as _g
from docx import Document

_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
BOOK = sorted(_g.glob(_ROOT + '/90_通用工具_全学期复用/90-A1_学生记录册*.docx'))[-1]
_CN = '一二三四五六七八九'
MAIN = {}
for _d in sorted(_g.glob(_ROOT + '/0*_项目*')):
    _m = re.search(r'/0(\d)_项目', _d.replace('\\', '/'))
    if not _m:
        continue
    _f = sorted(_g.glob(_d + '/0*主文档*.docx'))
    if _f:
        MAIN[_CN[int(_m.group(1)) - 1]] = _f[-1]

def allp(doc):
    for p in doc.paragraphs: yield p
    for tb in doc.tables:
        for r in tb.rows:
            for c in r.cells:
                for p in c.paragraphs: yield p

def main_pauses(path, cn):
    d = Document(path)
    out = {}
    for p in d.paragraphs:
        t = p.text.strip()
        if t.startswith('■'):
            m = re.search(r'项目%s第(\d+)暂停点' % cn, t)
            if m:
                q = re.sub(r'^■\s*项目.第\d+暂停点\s*　?请暂停\d+秒。', '', t).strip()
                out[int(m.group(1))] = q
    return out

def book_pauses(cn):
    d = Document(BOOK)
    out = {}
    for p in allp(d):
        t = p.text.strip()
        m = re.search(r'项目%s第(\d+)暂停点' % cn, t)
        if m:
            q = re.sub(r'^项目.第\d+暂停点\s*　?', '', t).strip()
            out[int(m.group(1))] = q
    # 五行表格里的「第N题」标注
    for tb in d.tables:
        if tb.rows[0].cells[0].text.strip() == '第几天':
            for r in tb.rows[1:]:
                m = re.search(r'第(\d+)题', r.cells[0].text)
                if m: out.setdefault(int(m.group(1)), '（五行表格内）'+r.cells[0].text.strip().replace('\n','/'))
    return out

bad = 0
for cn, path in MAIN.items():
    mp = main_pauses(path, cn)
    bp = book_pauses(cn)
    print('=' * 60)
    print('项目%s　主文档 %d 个暂停点　记录册 %d 个' % (cn, len(mp), len(bp)))
    miss = sorted(set(mp) - set(bp))
    extra = sorted(set(bp) - set(mp))
    if miss:
        print('  ★记录册缺:', miss); bad += 1
    if extra:
        print('  ★记录册多出:', extra); bad += 1
    if not miss and not extra:
        print('  编号完全对应 ✓')
    # 逐条比对题目意思（取前12字）
    diff = []
    for k in sorted(set(mp) & set(bp)):
        clean=lambda x: set(re.sub(r'[：:，。？?（）()「」★\s]|册子|请暂停\d+秒|写在|上写|那题', '', x))
        A,B=clean(mp[k]),clean(bp[k])
        share=len(A&B)/max(1,min(len(A),len(B)))
        if share < 0.5:
            diff.append((k, mp[k][:34], bp[k][:34]))
    if diff:
        print('  ★题目内容对不上的:')
        for k, a, b in diff:
            print('    第%-2d题  主文档:%s' % (k, a))
            print('           记录册:%s' % b)
        bad += 1
print('=' * 60)
print('结论:', '★有不一致，需修正' if bad else '全部一致 ✓')
